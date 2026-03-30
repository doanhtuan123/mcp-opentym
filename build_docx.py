# -*- coding: utf-8 -*-
"""
MCP_GUIDE.md  →  MCP_GUIDE.docx
Dùng Word built-in styles + python-docx.
Font: Times New Roman (body/heading), Courier New (code)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = Path(__file__).parent / "MCP_GUIDE.md"
DEST = Path(__file__).parent / "MCP_GUIDE.docx"

FONT_BODY = "Times New Roman"
FONT_CODE = "Courier New"

# ─────────────────────────────────────────────────────────
# PHẦN 1: Thiết lập document — page margin, font mặc định
# ─────────────────────────────────────────────────────────

def setup_document(doc: Document):
    """Margin trang + font mặc định toàn document."""
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Font mặc định
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.line_spacing = Pt(18)
    _set_east_asian_font(normal.element, FONT_BODY)


def _set_east_asian_font(element, font_name: str):
    """Gắn font East Asian để hiển thị đúng tiếng Việt."""
    rFonts = element.find(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    if rFonts is None:
        return
    rFonts.set(qn("w:eastAsia"), font_name)


# ─────────────────────────────────────────────────────────
# PHẦN 2: Thiết lập Heading styles (H1–H4)
# ─────────────────────────────────────────────────────────

HEADING_CONFIG = {
    1: {"size": 20, "color": RGBColor(0x1F, 0x49, 0x7D), "space_before": 18, "border": True},
    2: {"size": 16, "color": RGBColor(0x2E, 0x74, 0xB5), "space_before": 14, "border": False},
    3: {"size": 13, "color": RGBColor(0x20, 0x96, 0xD6), "space_before": 10, "border": False},
    4: {"size": 12, "color": RGBColor(0x2E, 0x74, 0xB5), "space_before": 8,  "border": False},
}

def setup_heading_styles(doc: Document):
    """Override Heading 1–4 styles với font/màu/size mong muốn."""
    for level, cfg in HEADING_CONFIG.items():
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        style.font.name      = FONT_BODY
        style.font.size      = Pt(cfg["size"])
        style.font.bold      = True
        style.font.color.rgb = cfg["color"]
        style.font.italic    = False
        style.paragraph_format.space_before   = Pt(cfg["space_before"])
        style.paragraph_format.space_after    = Pt(6)
        style.paragraph_format.keep_with_next = True

        # East Asian font
        rPr = style._element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
        )
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), FONT_BODY)


# ─────────────────────────────────────────────────────────
# PHẦN 3: Helper — gắn font cho một run
# ─────────────────────────────────────────────────────────

def apply_run_font(run, font_name: str, size: float,
                   bold=False, italic=False, color: RGBColor = None):
    """Gắn font name/size/style vào một run, kể cả East Asian."""
    run.font.name  = font_name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


# ─────────────────────────────────────────────────────────
# PHẦN 4: Parse inline markdown trong một dòng text
#          **bold**, *italic*, `code`, text thường
# ─────────────────────────────────────────────────────────

INLINE_PATTERN = re.compile(r'`([^`]+)`|\*\*(.+?)\*\*|\*(.+?)\*')

def render_inline(para, text: str, base_size: float = 12, base_bold: bool = False):
    """Thêm các run inline vào paragraph, xử lý **bold** / *italic* / `code`."""
    last = 0
    for m in INLINE_PATTERN.finditer(text):
        # Text thường trước match
        if m.start() > last:
            run = para.add_run(text[last:m.start()])
            apply_run_font(run, FONT_BODY, base_size, bold=base_bold)

        if m.group(1) is not None:       # `inline code`
            run = para.add_run(m.group(1))
            apply_run_font(run, FONT_CODE, 10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

        elif m.group(2) is not None:     # **bold**
            run = para.add_run(m.group(2))
            apply_run_font(run, FONT_BODY, base_size, bold=True)

        elif m.group(3) is not None:     # *italic*
            run = para.add_run(m.group(3))
            apply_run_font(run, FONT_BODY, base_size, italic=True)

        last = m.end()

    # Phần còn lại sau match cuối
    if last < len(text):
        run = para.add_run(text[last:])
        apply_run_font(run, FONT_BODY, base_size, bold=base_bold)


# ─────────────────────────────────────────────────────────
# PHẦN 5: Thêm Heading — dùng Word style "Heading N"
# ─────────────────────────────────────────────────────────

def add_heading(doc: Document, raw_text: str, level: int):
    # Xoá anchor {#...} và markdown link [text](url)
    text = re.sub(r'\{#[^}]+\}', '', raw_text).strip()
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    para = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    para.clear()  # xoá run mặc định

    run = para.add_run(text)
    cfg = HEADING_CONFIG.get(level, HEADING_CONFIG[4])
    apply_run_font(run, FONT_BODY, cfg["size"], bold=True, color=cfg["color"])

    # H1: thêm đường kẻ dưới
    if level == 1:
        _add_bottom_border(para, color="1F497D")


def _add_bottom_border(para, color: str):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────
# PHẦN 6: Code block — dùng bảng 1 ô nền xám
# ─────────────────────────────────────────────────────────

def add_code_block(doc: Document, code_lines: list[str]):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]

    # Tô nền xám cho ô
    _shade_cell(cell, "F2F2F2")

    # Xoá paragraph mặc định trong cell
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    # Thêm từng dòng code
    for line in code_lines:
        para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing = Pt(14)

        run = para.add_run(line or " ")
        apply_run_font(run, FONT_CODE, 9.5, color=RGBColor(0x1F, 0x1F, 0x1F))

    # Khoảng trắng sau bảng
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _shade_cell(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


# ─────────────────────────────────────────────────────────
# PHẦN 7: Bảng Markdown
# ─────────────────────────────────────────────────────────

def add_markdown_table(doc: Document, raw_rows: list[str]):
    # Parse từng hàng thành list cells
    rows = [
        [c.strip() for c in r.strip().strip("|").split("|")]
        for r in raw_rows
    ]
    # Bỏ dòng separator (---|---)
    rows = [r for r in rows if not all(re.match(r'^-+$', c.replace(":", "")) for c in r if c)]
    if not rows:
        return

    ncols = max(len(r) for r in rows)
    tbl   = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"

    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell  = tbl.cell(ri, ci)
            value = row[ci] if ci < len(row) else ""

            if ri == 0:
                # Header row: nền xanh, chữ trắng đậm
                _shade_cell(cell, "2E74B5")
                run = cell.paragraphs[0].add_run(
                    re.sub(r'[`*]', '', value)  # bỏ markdown trong header
                )
                apply_run_font(run, FONT_BODY, 11, bold=True,
                               color=RGBColor(0xFF, 0xFF, 0xFF))
            else:
                # Data row: xen kẽ trắng/xanh nhạt
                _shade_cell(cell, "FFFFFF" if ri % 2 == 1 else "EAF1FB")
                render_inline(cell.paragraphs[0], value, base_size=11)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─────────────────────────────────────────────────────────
# PHẦN 8: List item (bullet / numbered)
# ─────────────────────────────────────────────────────────

def add_list_item(doc: Document, text: str, level: int = 0,
                  ordered: bool = False, counter: int = 1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before      = Pt(1)
    para.paragraph_format.space_after       = Pt(2)
    para.paragraph_format.left_indent       = Inches(0.25 * (level + 1))
    para.paragraph_format.first_line_indent = Inches(-0.25)

    marker = f"{counter}." if ordered else "•"
    run_marker = para.add_run(f"{marker}  ")
    apply_run_font(run_marker, FONT_BODY, 12, bold=True,
                   color=RGBColor(0x2E, 0x74, 0xB5))

    render_inline(para, text, base_size=12)


# ─────────────────────────────────────────────────────────
# PHẦN 9: Blockquote / Note
# ─────────────────────────────────────────────────────────

def add_blockquote(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.4)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)

    # Viền trái xanh
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "2E74B5")
    pBdr.append(left)
    pPr.append(pBdr)

    # Nền xanh nhạt
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EAF4FF")
    pPr.append(shd)

    render_inline(para, text, base_size=11)


# ─────────────────────────────────────────────────────────
# PHẦN 10: Horizontal rule
# ─────────────────────────────────────────────────────────

def add_hr(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)

    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BBBBBB")
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────
# PHẦN 11: Normal paragraph
# ─────────────────────────────────────────────────────────

def add_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    render_inline(para, text, base_size=12)


# ─────────────────────────────────────────────────────────
# PHẦN 12: MAIN PARSER — đọc markdown, dispatch từng block
# ─────────────────────────────────────────────────────────

def convert(src: Path, dest: Path):
    doc = Document()
    setup_document(doc)
    setup_heading_styles(doc)

    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Code block (``` ... ```) ──────────────────────
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # bỏ dòng closing ```
            continue

        # ── Horizontal rule (--- hoặc ***) ───────────────
        if re.match(r'^[-*]{3,}$', line.strip()):
            add_hr(doc)
            i += 1
            continue

        # ── Heading (# / ## / ### / ####) ────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            add_heading(doc, m.group(2), level=len(m.group(1)))
            i += 1
            continue

        # ── Blockquote (> ...) ───────────────────────────
        if line.startswith(">"):
            add_blockquote(doc, re.sub(r'^>\s*', '', line))
            i += 1
            continue

        # ── Markdown table (| ... |) ─────────────────────
        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        # ── Numbered list (1. / 2. / ...) ────────────────
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m:
            level   = len(m.group(1)) // 2
            counter = int(m.group(2))
            add_list_item(doc, m.group(3), level=level, ordered=True, counter=counter)
            i += 1
            continue

        # ── Bullet list (- / * / +) ───────────────────────
        m = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if m:
            level = len(m.group(1)) // 2
            add_list_item(doc, m.group(2), level=level, ordered=False)
            i += 1
            continue

        # ── Dòng trống ───────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────
        add_paragraph(doc, line)
        i += 1

    doc.save(str(dest))
    print(f"[OK] {dest.name}  ({dest.stat().st_size // 1024} KB)")


# ─────────────────────────────────────────────────────────
if __name__ == "__main__":
    convert(SRC, DEST)
